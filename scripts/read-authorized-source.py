"""只读提取授权来源文本；输出前遮蔽常见敏感标识。"""
from __future__ import annotations

import re
import sys
from pathlib import Path


def scrub(text: str) -> str:
    text = re.sub(r"(?<!\d)1[3-9]\d{9}(?!\d)", "[手机号已遮蔽]", text)
    text = re.sub(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", "[邮箱已遮蔽]", text, flags=re.I)
    text = re.sub(r"(?<!\d)\d{17}[\dXx](?!\d)", "[身份证号已遮蔽]", text)
    return text


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    blocks = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    return "\n".join((page.extract_text() or "").strip() for page in PdfReader(path).pages)


def read_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[str] = []
    for sheet in workbook.worksheets:
        blocks.append(f"[工作表] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() if value is not None else "" for value in row]
            if any(values):
                blocks.append(" | ".join(values))
    return "\n".join(blocks)


def main() -> None:
    for raw in sys.argv[1:]:
        path = Path(raw)
        suffix = path.suffix.lower()
        if suffix == ".docx":
            text = read_docx(path)
        elif suffix == ".pdf":
            text = read_pdf(path)
        elif suffix == ".xlsx":
            text = read_xlsx(path)
        else:
            raise ValueError(f"不支持的格式：{path}")
        print(f"\n===== {path.name} =====")
        print(scrub(text))


if __name__ == "__main__":
    main()
